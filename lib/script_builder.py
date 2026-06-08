"""SCOTT Automation Deck Builder — selling script builder.

Generates a Word doc that accompanies the pitch deck. For each Why slide it
gives the rep: what to say, what NOT to say, discovery questions, and a
transition line. Plus a StoryBrand cheat sheet and objection responses.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .stats_library import get_stat


# Color tokens matching the deck
SCOTT_BLUE = RGBColor(0x00, 0x46, 0x7F)
ACCENT_SLATE = RGBColor(0x74, 0x8F, 0xA6)
DARK = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x5F, 0x5E, 0x5A)


# ============================================================================
# Document setup
# ============================================================================

def _setup_document() -> Document:
    """Create a Document with our brand styles configured."""
    doc = Document()

    # Page setup — letter size, 0.75" margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Default body style
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK

    # Customize heading 1
    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = SCOTT_BLUE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    # Customize heading 2
    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = SCOTT_BLUE
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    return doc


# ============================================================================
# Paragraph helpers
# ============================================================================

def _add_eyebrow(doc: Document, text: str):
    """Small uppercase slate-blue label."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = ACCENT_SLATE
    # Slight letter spacing for eyebrow feel
    rPr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "30")
    rPr.append(spacing)


def _add_body(doc: Document, text: str, *, italic: bool = False, color: RGBColor | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.color.rgb = color or DARK
    return p


def _add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK


def _add_quote_line(doc: Document, text: str):
    """Italic indented line for transition quotes."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"\u201C{text}\u201D")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = DARK


def _add_divider(doc: Document):
    """Thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D5D1C7")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================================================================
# Slide card — reusable block for each slide in the deck
# ============================================================================

def _add_slide_card(doc, *, slide_label, slide_title, runtime, talking_points,
                    do_not_say, discovery_qs, transition_line):
    """One slide's worth of talking guidance."""
    h2 = doc.add_heading(f"{slide_label} \u2014 {slide_title}", level=2)
    h2.paragraph_format.space_before = Pt(14)

    if runtime:
        _add_body(doc, f"Suggested time: {runtime}", italic=True, color=MUTED)

    _add_eyebrow(doc, "What to say")
    for line in talking_points:
        _add_bullet(doc, line)

    _add_eyebrow(doc, "What NOT to say")
    for line in do_not_say:
        _add_bullet(doc, line)

    if discovery_qs:
        _add_eyebrow(doc, "Discovery questions to ask the customer")
        for q in discovery_qs:
            _add_bullet(doc, q)

    _add_eyebrow(doc, "Transition to the next slide")
    _add_quote_line(doc, transition_line)

    _add_divider(doc)


# ============================================================================
# Slide-card content generators (parameterized by the rep's form input)
# ============================================================================

def _pain_point_card(doc, pain_point: str, customer_name: str):
    """Slide 4 — "Here's what we heard" — read the customer's own words back to them."""
    excerpt = (pain_point[:120] + "…") if len(pain_point) > 120 else pain_point
    _add_slide_card(
        doc,
        slide_label="Slide 4",
        slide_title="Here's what we heard",
        runtime="2\u20133 minutes",
        talking_points=[
            "Read the slide back to them, slowly, in their own words. Pause after each clause.",
            "Make eye contact. Don't talk over the slide \u2014 let them see their own situation in print.",
            f"Say something like: \u201CThis is what we heard from your team. Did we get it right? Anything missing?\u201D",
            "If they correct you, write it down. Then say: \u201CThank you, that's important \u2014 I want to make sure we're solving the right problem.\u201D",
        ],
        do_not_say=[
            "Anything starting with \u201CSCOTT has been doing this for 75 years.\u201D Save your r\u00e9sum\u00e9. We're not there yet.",
            "Any feature, product, or capability. They don't care yet.",
            "\u201CWe can fix that\u201D \u2014 too soon. You haven't earned that statement.",
        ],
        discovery_qs=[
            "Is there something I missed that's making this worse than I just described?",
            f"Who else at {customer_name or 'your company'} is feeling this \u2014 operations, HR, finance, all of the above?",
            "How long has it been like this?",
        ],
        transition_line=(
            "What you're dealing with isn't unique \u2014 and the national picture tells us "
            "why this kind of problem is getting harder to ignore."
        ),
    )


def _stat_card(doc, stat: dict, slide_num: int, position: str):
    """A 'national picture' stat slide. `position` is e.g. 'Slide 5'."""
    _add_slide_card(
        doc,
        slide_label=position,
        slide_title=stat["headline"].title(),
        runtime="2 minutes",
        talking_points=[
            f"Lead with the big number out loud: \u201C{stat['big_number']}.\u201D Let it sit for a second.",
            f"Then read the supporting line. The number does the work \u2014 don't editorialize.",
            "Bridge it to the customer: \u201CYou're not crazy for being worried about this. The data says you're paying attention to the right thing.\u201D",
            "Cite the source if asked, but don't lead with it. Credibility comes from being calm, not from name-dropping.",
        ],
        do_not_say=[
            "Don't compare them to their competitors \u2014 comparisons make customers defensive.",
            "Don't follow this slide with a SCOTT pitch. Stay in their world for one more beat.",
            "Don't tell them how scary the number is. The number is scary on its own.",
        ],
        discovery_qs=[
            "How does your performance compare to that industry benchmark?",
            "What's one of these costing you today, in dollars or in lost hours?",
        ],
        transition_line=(
            "When you can see the cost of staying the same, the case for changing something "
            "gets a lot clearer. Let me show you what changing it actually looks like."
        ),
    )


def _success_state_card(doc, success_state: str):
    excerpt = success_state[:100] + ("…" if len(success_state) > 100 else "")
    _add_slide_card(
        doc,
        slide_label="What you get",
        slide_title="The transformation",
        runtime="2 minutes",
        talking_points=[
            "Read each clause separately. Don't rush.",
            "After each promise, pause. Let them picture it.",
            "Close with a line that connects the change to their people, not just their numbers \u2014 promotion, growth, easier work.",
            "This is the StoryBrand emotional close. Their team isn't replaced; their team is elevated.",
        ],
        do_not_say=[
            "Don't promise specific dollar savings unless they came from a real engineering estimate.",
            "Don't say \u201Cthis will solve all your problems.\u201D Hyperbole erodes trust.",
            "Don't say \u201Crobots replace workers.\u201D Say \u201Crobots free workers for better work.\u201D",
        ],
        discovery_qs=[
            "If this future were real next quarter, what would you do with the time and budget it freed up?",
            "Who on your team would be the natural fit for the new roles automation creates?",
        ],
        transition_line=(
            "That's the destination. Now let me show you the route \u2014 and how short it actually is."
        ),
    )


def _plan_card(doc, steps):
    """Plan slide — three steps."""
    _add_slide_card(
        doc,
        slide_label="The plan",
        slide_title="How we get you there",
        runtime="3\u20134 minutes",
        talking_points=[
            "Walk all three steps in order. Don't skip ahead.",
            f"On step 1 ({steps[0].get('title', '')}): emphasize the deliverable, not the activity. Customers care about what they'll see.",
            f"On step 2 ({steps[1].get('title', '')}): emphasize that their line keeps running. Disruption is one of their biggest fears \u2014 address it directly.",
            f"On step 3 ({steps[2].get('title', '')}): emphasize ownership. Their team runs it after we hand it off; we're not creating a dependency.",
            "Close with: \u201CThe reason this is a 3-step plan is that we've done it dozens of times. Each step has a clear deliverable and a clear handoff.\u201D",
        ],
        do_not_say=[
            "Don't list more than 3 steps. Three is the trustable number; 5+ feels overwhelming.",
            "Don't get into hardware specs (robot model, PLC brand). That's the next conversation.",
            "Don't say \u201Cit depends\u201D without immediately following up with what it depends on.",
        ],
        discovery_qs=[
            "Which of the three steps would you want to see firsthand?",
            "Who on your team should be involved at each step?",
        ],
        transition_line="The plan is short on purpose. The only thing left is choosing when to start.",
    )


def _cta_card(doc, cta_text: str):
    _add_slide_card(
        doc,
        slide_label="The call to action",
        slide_title="Your next step",
        runtime="1\u20132 minutes",
        talking_points=[
            "Slow down. This is the most important slide.",
            "Read the offer exactly as it appears on the slide. Don't paraphrase.",
            "Then ask, directly: \u201CDoes that work, or would the week after be better?\u201D",
            "Notice the question \u2014 you're not asking \u201Cyes or no.\u201D You're asking which time. That's a StoryBrand assumptive close.",
            "If they hesitate, the script is: \u201CWhat would you need to see between now and then to feel good about that meeting?\u201D That uncovers the real objection.",
        ],
        do_not_say=[
            "Don't ask \u201CWhat do you think?\u201D \u2014 too open-ended, lets them off the hook.",
            "Don't offer a discount. You haven't quoted anything yet.",
            "Don't follow this slide with a long features dump. The decision moment is here.",
        ],
        discovery_qs=[
            "Who else needs to be at the next meeting \u2014 operations, EHS, finance?",
            "Is there anything specific you'd want our engineering team to look at first?",
        ],
        transition_line=(
            "With that on the calendar, let me give you the rest of the picture on who SCOTT "
            "is and the kind of work we do \u2014 so you know exactly who's showing up next."
        ),
    )


# ============================================================================
# Top-level build function
# ============================================================================

def build_script(form_data: dict, output_path: str):
    """Generate a Word doc selling script at output_path from form data."""
    doc = _setup_document()

    customer_name = form_data.get("customerName", "").strip() or "the customer"
    pain_point = form_data.get("painPoint", "").strip()
    success_state = form_data.get("successState", "").strip()
    plan = form_data.get("plan", {})
    cta = form_data.get("callToAction", "").strip()
    value_drivers = form_data.get("valueDrivers", [])

    # === Cover ===
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("PITCH DECK SELLING SCRIPT")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = ACCENT_SLATE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Help the customer be the hero")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = SCOTT_BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(f"Customer: {customer_name}")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = DARK

    _add_body(doc, (
        f"This script accompanies the {customer_name} pitch deck. It uses the StoryBrand method: "
        "the customer is the hero of the story; SCOTT is the guide who hands them a plan. Each "
        "slide below tells you what to say, what to avoid, what to ask, and how to bridge to the "
        "next slide. Read it once before the meeting. Don't read it during the meeting."
    ))

    # === Before you walk in ===
    _add_eyebrow(doc, "Before you walk in")
    _add_bullet(doc, "Re-read the customer's pain point in your own words. If you can't say it back to them better than they said it to you, you're not ready.")
    _add_bullet(doc, "Know your stat slides cold. Each big number is doing work for you \u2014 say them slowly.")
    _add_bullet(doc, "Your job in the first 10 minutes is to make them feel heard, not to talk about SCOTT. They will care about SCOTT only after they feel understood.")
    _add_bullet(doc, "Bring a pen and pad. Take a note while they're talking on slide 4. They'll see you doing it. It signals you're a guide, not a vendor.")

    # === Part 1: The problem ===
    if pain_point:
        doc.add_heading("Part 1: The problem", level=1)
        _add_body(doc, "Goal: Make the customer feel that you understand them \u2014 and that they're not alone. Do not pitch anything yet.")
        _pain_point_card(doc, pain_point, customer_name)
        # Stat cards
        for i, tag in enumerate(value_drivers):
            stat = get_stat(tag)
            if not stat:
                continue
            position = f"Stat slide ({stat['big_number']})"
            _stat_card(doc, stat, slide_num=5 + i, position=position)

    # === Part 2: The transformation ===
    if success_state:
        doc.add_heading("Part 2: The transformation", level=1)
        _add_body(doc, "Goal: Paint the after-picture. Make them feel what success looks like before you explain how to get there.")
        _success_state_card(doc, success_state)

    # === Part 3: The plan ===
    has_any_step = any(plan.get(f"step{i}", {}).get("title") for i in (1, 2, 3))
    if has_any_step:
        doc.add_heading("Part 3: The plan", level=1)
        _add_body(doc, "Goal: Make the path forward feel concrete, simple, and safe. Three steps \u2014 no more.")
        steps = [plan.get(f"step{i}", {}) for i in (1, 2, 3)]
        _plan_card(doc, steps)

    # === Part 4: The call to action ===
    if cta:
        doc.add_heading("Part 4: The call to action", level=1)
        _add_body(doc, "Goal: Give them one small, low-friction yes. Not the buying decision yet \u2014 the next-step decision.")
        _cta_card(doc, cta)

    # === Appendix: StoryBrand cheat sheet ===
    doc.add_heading("Appendix A: StoryBrand cheat sheet", level=1)
    _add_body(doc, "If you forget everything else in this script, remember the seven beats:")
    _add_bullet(doc, "A hero \u2014 the customer. Not SCOTT.")
    _add_bullet(doc, "Has a problem \u2014 name it in their language, not yours.")
    _add_bullet(doc, "Meets a guide \u2014 that's SCOTT. Guides have empathy and authority. Show empathy by listening; show authority by referencing data and past work.")
    _add_bullet(doc, "Who gives them a plan \u2014 three steps. Always three.")
    _add_bullet(doc, "And calls them to action \u2014 a single, specific, time-bound next step.")
    _add_bullet(doc, "So they avoid failure \u2014 name what stays the same if they don't act.")
    _add_bullet(doc, "And achieve success \u2014 the after-picture. Make it specific and emotional.")
    _add_body(doc, "If you find yourself talking about SCOTT for more than 30 seconds before you've talked about them, stop. Start over. Their story is the story.", italic=True, color=MUTED)

    # === Appendix B: Objection responses ===
    doc.add_heading("Appendix B: Objection responses", level=1)

    doc.add_heading("\u201CWe can't afford this right now.\u201D", level=2)
    _add_body(doc, (
        "Acknowledge it. Then: \u201CWhat you can't afford is staying with the status quo for another "
        "year. The walkthrough is free \u2014 let's at least put a real number on what the change "
        "would cost so you can compare it to the cost of not changing.\u201D"
    ))

    doc.add_heading("\u201COur team won't accept automation.\u201D", level=2)
    _add_body(doc, (
        "\u201CThat comes up a lot, and it's almost always a sign that the rollout wasn't done right. "
        "The reason we put operator training in step 3 and not as an afterthought is that the people "
        "running the line need to feel like the cell was built around them, not against them. We'd "
        "want your most skeptical operator at the walkthrough.\u201D"
    ))

    doc.add_heading("\u201CHow long does this take?\u201D", level=2)
    _add_body(doc, (
        "\u201CFor a project of this scope, typically six to nine months from design kickoff to "
        "production. The walkthrough we're proposing is what gets the clock started.\u201D"
    ))

    doc.add_heading("\u201CWhy SCOTT and not a competitor?\u201D", level=2)
    _add_body(doc, (
        "\u201CFair question. Three things: we design and build in-house \u2014 same team end-to-end, "
        "so there's no finger-pointing if something goes wrong. We have a 75-year service record in "
        "the region, so if you call us in three years for a question, the person who built it picks "
        "up. And we're an integrator, not a robot reseller \u2014 we don't push a specific brand if "
        "it's wrong for your application. You can hear all of that in the rest of the deck, but those "
        "are the headlines.\u201D"
    ))

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run("\u2014")
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED

    _add_body(doc, (
        "This script was generated alongside your pitch deck. Update the customer pain point, value "
        "drivers, plan, and call-to-action in the deck builder, and the script will regenerate to match."
    ), italic=True, color=MUTED)

    doc.save(output_path)
